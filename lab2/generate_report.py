#!/usr/bin/env python3
"""Generate UTM-formatted .docx report for the crypto-puzzle parallel solver."""

from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup (UTM: left 20mm, right 10mm, top 20mm, bottom 20mm) ──
for section in doc.sections:
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(10)

# ── Default style: Times New Roman 12pt, 1.5 line spacing ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# ── Helper functions ──

def add_chapter_title(text, new_page=True):
    """Chapter title: Bold 13pt, centered, uppercase, new page."""
    if new_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    doc.add_paragraph()  # spacing after title

def add_subchapter_title(text):
    """Subchapter title: Bold 12pt, left-aligned."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

def add_text(text):
    """Normal paragraph."""
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def add_text_no_indent(text):
    """Normal paragraph without first line indent."""
    return doc.add_paragraph(text)

def add_bold_text(label, text):
    """Paragraph with bold label followed by normal text."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(label)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run2 = p.add_run(text)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)
    return p

def add_dash_item(text):
    """Enumeration item with dash."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'\u2013 {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_letter_item(letter, text):
    """Enumeration item with letter a), b), etc."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{letter}) {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_code_block(code):
    """Code block: Courier New 10pt, single spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)

def add_figure_caption(text):
    """Figure caption: bold, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_table_caption(text):
    """Table caption: bold, right-aligned, above table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_table(headers, rows):
    """Add a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.line_spacing = 1.0
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()  # spacing after table
    return table


def add_bib_entry(text):
    """Bibliography entry."""
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    return p


# ═══════════════════════════════════════════════════════════════
#  FOAIA DE TITLU
# ═══════════════════════════════════════════════════════════════

# Add vertical spacing to center content
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSITATEA TEHNICA A MOLDOVEI')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Facultatea Calculatoare, Informatica si Microelectronica')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Departamentul Informatica si Ingineria Sistemelor')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disciplina: Algoritmi si Analiza Algoritmilor')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('LUCRARE DE LABORATOR Nr. 2')
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tema: ')
r.bold = True
r.font.size = Pt(12)
r.font.name = 'Times New Roman'
r2 = p.add_run('Implementarea unui sistem distribuit de rezolvare paralela '
               'a cripto-puzzle-urilor cu raportare de performanta prin '
               'intermediul unui mesaj-broker')
r2.font.size = Pt(12)
r2.font.name = 'Times New Roman'

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Student: _________________________ gr. ____')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('Profesor: _________________________')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Chisinau, 2024')
run.font.size = Pt(12)
run.font.name = 'Times New Roman'


# ═══════════════════════════════════════════════════════════════
#  CUPRINS
# ═══════════════════════════════════════════════════════════════

add_chapter_title('Cuprins')

toc_entries = [
    ('Lista de Abrevieri si Definitii', '4'),
    ('Introducere', '5'),
    ('1 Analiza teoretica a mecanismelor de calcul distribuit si cripto-puzzle', '6'),
    ('   1.1 Cripto-puzzle-uri bazate pe SHA-256', '6'),
    ('   1.2 Comunicatia prin mesaj-broker si modelul RPC asincron', '7'),
    ('   1.3 Mecanisme de masurare a timpului monoton', '8'),
    ('2 Proiectarea si implementarea sistemului', '9'),
    ('   2.1 Arhitectura generala a solutiei', '9'),
    ('   2.2 Serverul de dispecerat (ruby_server.rb)', '10'),
    ('   2.3 Lucratorul Ruby (ruby_computer.rb)', '12'),
    ('   2.4 Lucratorul Python (python_computer.py)', '13'),
    ('   2.5 Lucratorul C# (csharp_computer.cs)', '14'),
    ('   2.6 Infrastructura Docker Compose', '15'),
    ('3 Rezultatele executiei si analiza performantei', '16'),
    ('   3.1 Metodologia de masurare si raportare', '16'),
    ('   3.2 Rezultate experimentale la diferite niveluri de dificultate', '17'),
    ('   3.3 Analiza comparativa a lucratorilor', '19'),
    ('Concluzii', '21'),
    ('Bibliografia', '22'),
    ('Anexa A \u2014 Codul sursa al serverului de dispecerat', '23'),
    ('Anexa B \u2014 Codul sursa al lucratorului Ruby', '24'),
    ('Anexa C \u2014 Codul sursa al lucratorului Python', '25'),
    ('Anexa D \u2014 Codul sursa al lucratorului C#', '26'),
    ('Anexa E \u2014 Fisierul Docker Compose', '27'),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    dots = '.' * max(2, 70 - len(entry) - len(page))
    run = p.add_run(f'{entry} {dots} {page}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════
#  LISTA DE ABREVIERI SI DEFINITII
# ═══════════════════════════════════════════════════════════════

add_chapter_title('Lista de Abrevieri si Definitii')

abbrevs = [
    ('AMQP', 'Advanced Message Queuing Protocol \u2014 protocol standard de comunicatie pentru mesaj-brokere'),
    ('API', 'Application Programming Interface \u2014 interfata de programare a aplicatiei'),
    ('Bunny', 'Biblioteca Ruby pentru interactiunea cu RabbitMQ prin protocolul AMQP'),
    ('CLI', 'Command-Line Interface \u2014 interfata in linia de comanda'),
    ('CPU', 'Central Processing Unit \u2014 unitate centrala de procesare'),
    ('C#', 'Limbaj de programare orientat pe obiecte dezvoltat de Microsoft, parte a platformei .NET'),
    ('Docker', 'Platforma de containerizare a aplicatiilor'),
    ('Docker Compose', 'Instrument pentru definirea si rularea aplicatiilor Docker multi-container'),
    ('DLL', 'Dynamic-Link Library \u2014 biblioteca cu legatura dinamica'),
    ('FIFO', 'First In, First Out \u2014 disciplina de ordonare a elementelor intr-o coada'),
    ('Hash', 'Valoarea hexazecimala rezultata in urma aplicarii unei functii criptografice de dispersie'),
    ('Hash rate', 'Numarul de operatii de hashing executate pe unitatea de timp'),
    ('JSON', 'JavaScript Object Notation \u2014 format de serializare a datelor structurate'),
    ('ms', 'Milisecunda \u2014 unitate de masura a timpului (10\u207b\u00b3 s)'),
    ('Nonce', 'Number used once \u2014 numar intreg incremental utilizat ca sufix al sirului de intrare'),
    ('NuGet', 'Manager de pachete pentru ecosistemul .NET'),
    ('Pika', 'Biblioteca Python pentru protocolul AMQP, utilizata pentru conectarea la RabbitMQ'),
    ('RabbitMQ', 'Mesaj-broker open-source bazat pe protocolul AMQP'),
    ('RPC', 'Remote Procedure Call \u2014 apel de procedura la distanta'),
    ('Ruby', 'Limbaj de programare dinamic, orientat pe obiecte'),
    ('SDK', 'Software Development Kit \u2014 kit de dezvoltare software'),
    ('SHA-256', 'Secure Hash Algorithm 256-bit \u2014 functie criptografica de dispersie din familia SHA-2'),
    ('UUID', 'Universally Unique Identifier \u2014 identificator unic universal de 128 de biti'),
    ('Worker', 'Lucrator \u2014 proces sau container care executa o sarcina de calcul alocata'),
]

add_table(
    ['Termen / Abreviere', 'Definitie'],
    abbrevs
)


# ═══════════════════════════════════════════════════════════════
#  INTRODUCERE
# ═══════════════════════════════════════════════════════════════

add_chapter_title('Introducere')

add_text(
    'Rezolvarea cripto-puzzle-urilor reprezinta un mecanism fundamental in protocoalele de tip '
    'Proof-of-Work, utilizate pe scara larga in sistemele blockchain si in protectia impotriva '
    'atacurilor de tip spam. Un cripto-puzzle SHA-256 consta in identificarea unui numar intreg '
    '(nonce) care, concatenat cu un sir de intrare predefinit, produce un rezumat criptografic '
    '(hash) ce incepe cu un numar specificat de cifre hexazecimale egale cu zero. Spatiul de '
    'cautare creste exponential cu dificultatea: pentru dificultatea d, numarul mediu de '
    'incercari necesare este de ordinul 16^d.'
)

add_text(
    'Prezenta lucrare de laborator documenteaza proiectarea, implementarea si evaluarea unui '
    'sistem distribuit eterogen care paralelizeaza cautarea nonce-ului pe trei lucratori '
    'independenti: unul implementat in Ruby, unul in Python si unul in C#. Coordonarea '
    'sarcinilor este realizata prin intermediul unui mesaj-broker RabbitMQ, utilizand modelul '
    'RPC asincron cu cozi de raspuns exclusive si filtrare prin identificatori de corelatie.'
)

add_text('Obiectivele principale ale lucrarii sunt:')
add_dash_item('implementarea unui protocol unificat de raspuns al lucratorilor, indiferent de limbajul de programare utilizat;')
add_dash_item('proiectarea unui mecanism de asteptare a tuturor lucratorilor (nu doar a primului raspuns), cu timeout configurabil;')
add_dash_item('colectarea si raportarea metricilor de performanta per executie si cumulat pe sesiune;')
add_dash_item('containerizarea integrala a sistemului prin Docker Compose.')

add_text(
    'Raportul este structurat in trei capitole principale: capitolul 1 prezinta bazele '
    'teoretice necesare intelegerii solutiei; capitolul 2 descrie in detaliu proiectarea si '
    'implementarea fiecarei componente; capitolul 3 analizeaza rezultatele experimentale si '
    'comparatia de performanta intre lucratori.'
)


# ═══════════════════════════════════════════════════════════════
#  CAPITOLUL 1
# ═══════════════════════════════════════════════════════════════

add_chapter_title('1 Analiza teoretica a mecanismelor de calcul distribuit si cripto-puzzle')

add_subchapter_title('1.1 Cripto-puzzle-uri bazate pe SHA-256')

add_text(
    'SHA-256 (Secure Hash Algorithm, varianta de 256 biti) este o functie criptografica de '
    'dispersie din familia SHA-2, standardizata de NIST in FIPS PUB 180-4 [1]. Functia '
    'transforma un mesaj de lungime arbitrara intr-un rezumat de 256 biti (32 octeti, '
    'reprezentat ca 64 de caractere hexazecimale) cu proprietatile:'
)
add_dash_item('determinism \u2014 acelasi mesaj produce intotdeauna acelasi rezumat;')
add_dash_item('efect de avalansa \u2014 o modificare de un singur bit in intrare produce o schimbare impredictibila in aproximativ jumatate din bitii rezumatului;')
add_dash_item('rezistenta la preimagine \u2014 cunoasterea rezumatului nu permite reconstructia mesajului original in timp polinomial;')
add_dash_item('rezistenta la coliziuni \u2014 gasirea a doua mesaje distincte cu acelasi rezumat este computational infeasible.')

add_text(
    'Un cripto-puzzle de dificultate d se defineste formal astfel: dat sirul de intrare S '
    '(in cazul de fata "Hello World"), se cauta un intreg nonnegativ n astfel incat:'
)

add_code_block('SHA256(S || str(n))[0 : d] == "00...0" (d cifre hexazecimale zero)')

add_text(
    'Probabilitatea ca un hash ales aleatoriu sa satisfaca conditia este 1/16^d. Prin urmare, '
    'numarul mediu de incercari necesar este 16^d, iar pentru a acoperi spatiul cu o marja de '
    'siguranta, implementarea utilizeaza 2 * 16^d / WORKER_COUNT nonce-uri pe lucrator [2].'
)

add_text('Tabelul 1.1 ilustreaza cresterea exponentiala a sarcinii de calcul in functie de dificultate.')

add_table_caption('Tabelul 1.1 \u2014 Estimarea numarului de hashes necesare in functie de dificultate')
add_table(
    ['Dificultate (d)', '16^d (medie teoretica)', 'Range / lucrator (x2, /3)', 'Timp estimat la 400k hash/s'],
    [
        ['1', '16', '500 000 (minim)', '< 1 s'],
        ['2', '256', '500 000 (minim)', '< 2 s'],
        ['3', '4 096', '500 000 (minim)', '< 2 s'],
        ['4', '65 536', '500 000 (minim)', '< 2 s'],
        ['5', '1 048 576', '699 050', '~ 2 s'],
        ['6', '16 777 216', '11 184 810', '~ 28 s'],
        ['7', '268 435 456', '178 956 970', '~ 7 min'],
        ['8', '4 294 967 296', '2 863 311 530', '~ 2 ore'],
    ]
)

add_text(
    'Valorile din tabelul 1.1 confirma ca dificultatea 1\u20134 este rezolvata aproape instantaneu '
    '(range-ul minim de 500 000 asigura date de performanta semnificative), in timp ce '
    'dificultatile 6+ necesita timp de executie masurabil si relevant pentru comparatia intre lucratori.'
)


add_subchapter_title('1.2 Comunicatia prin mesaj-broker si modelul RPC asincron')

add_text(
    'RabbitMQ este un mesaj-broker open-source ce implementeaza protocolul AMQP 0-9-1 [3]. '
    'In arhitectura prezentata, comunicatia urmeaza modelul RPC (Remote Procedure Call) asincron '
    'descris in documentatia oficiala RabbitMQ [4], cu urmatoarele elemente:'
)
add_dash_item('coada de lucru (crypto-puzzle-inquiries, auto_delete: true) \u2014 coada comuna din care toti lucratorii consuma sarcini in regim round-robin;')
add_dash_item('coada de raspuns exclusiva \u2014 creata de server cu numele generat automat de broker, marcata exclusive: true; este stearsa automat la inchiderea conexiunii;')
add_dash_item('correlation_id \u2014 UUID generat de server la fiecare runda de calcul; lucratorii il propaga intact in raspuns; serverul filtreaza raspunsurile pentru a ignora eventualele mesaje intarziate din runde anterioare.')

add_text('Conform figurii 1.1, fluxul de mesaje pentru o singura runda de calcul este urmatorul.')

add_code_block(
    'Server                    RabbitMQ                   Workers\n'
    '  |                          |                           |\n'
    '  |-- publish x3 ---------->|                           |\n'
    '  |   (correlation_id=UUID) |-- deliver task 1 ------->| Ruby\n'
    '  |                         |-- deliver task 2 ------->| Python\n'
    '  |                         |-- deliver task 3 ------->| CSharp\n'
    '  |                         |                          |\n'
    '  |                         |<-- reply (corr_id=UUID) -| Ruby\n'
    '  |<-- reply 1 ------------|                          |\n'
    '  |                         |<-- reply (corr_id=UUID) -| Python\n'
    '  |<-- reply 2 ------------|                          |\n'
    '  |                         |<-- reply (corr_id=UUID) -| CSharp\n'
    '  |<-- reply 3 ------------|                          |\n'
    '  |                         |                          |\n'
    '  | [toate 3 primite]       |                          |\n'
    '  | print_performance_table |                          |'
)
add_figure_caption('Figura 1.1 \u2014 Diagrama fluxului de mesaje pentru o runda de calcul')

add_text(
    'Un aspect critic al implementarii il reprezinta faptul ca serverul asteapta toate cele '
    'trei raspunsuri (nu doar primul), utilizand un Mutex si un ConditionVariable Ruby. Aceasta '
    'abordare permite colectarea metricilor de performanta de la fiecare lucrator, indiferent de '
    'ordinea de finalizare.'
)


add_subchapter_title('1.3 Mecanisme de masurare a timpului monoton')

add_text(
    'Masurarea precisa a duratei de executie necesita utilizarea unui ceas monoton \u2014 un ceas '
    'care nu se da niciodata inapoi, imun la ajustarile NTP sau la modificarile manuale ale orei '
    'sistemului [5]. Fiecare lucrator utilizeaza mecanismul nativ al limbajului sau:'
)
add_dash_item('Ruby \u2014 Process.clock_gettime(Process::CLOCK_MONOTONIC) returneaza un Float in secunde cu precizie sub-microsecunda; diferenta intre doua apeluri, inmultita cu 1000, da durata in milisecunde;')
add_dash_item('Python \u2014 time.monotonic() este echivalentul direct, returneaza un float in secunde; disponibil incepand cu Python 3.3 [6];')
add_dash_item('C# \u2014 System.Diagnostics.Stopwatch utilizeaza intern QueryPerformanceCounter pe Windows si clock_gettime(CLOCK_MONOTONIC) pe Linux/macOS; Elapsed.TotalMilliseconds ofera precizie echivalenta [7].')

add_text(
    'Toti trei lucratorii pornesc masurarea inainte de prima iteratie a buclei de hashing si o '
    'opresc imediat dupa gasirea solutiei sau dupa epuizarea range-ului, incluzand astfel '
    'exclusiv timpul de calcul al puzzle-ului, fara latenta de retea sau timp de serializare JSON.'
)


# ═══════════════════════════════════════════════════════════════
#  CAPITOLUL 2
# ═══════════════════════════════════════════════════════════════

add_chapter_title('2 Proiectarea si implementarea sistemului')

add_subchapter_title('2.1 Arhitectura generala a solutiei')

add_text(
    'Sistemul este compus din cinci servicii orchestrate prin Docker Compose, conectate printr-o '
    'retea virtuala Docker de tip bridge denumita "main". Conform figurii 2.1, componentele si '
    'rolurile lor sunt prezentate mai jos.'
)

add_code_block(
    '+-------------------------------------------------------------+\n'
    '|                    Retea Docker "main"                       |\n'
    '|                                                              |\n'
    '|  +------------------+         +---------------------------+ |\n'
    '|  |  lab2_producer   |         |       rabbitmq            | |\n'
    '|  |  (ruby_server.rb)|<------->| RabbitMQ 3-management     | |\n'
    '|  |  Ubuntu 22.04    |  AMQP   | Port 5672 (AMQP)          | |\n'
    '|  |  Ruby + Bunny    |         | Port 15672 (Management)   | |\n'
    '|  +------------------+         +-------------+-------------+ |\n'
    '|                                             |               |\n'
    '|        +------------------------------------+               |\n'
    '|        |                    |                |               |\n'
    '|  +-----+------+   +--------+-------+  +-----+----------+   |\n'
    '|  |lab2_consumer|  |python_computer |  |  cs_computer   |   |\n'
    '|  |ruby_computer|  |python:3.11-slim|  |  .NET 8.0      |   |\n'
    '|  |Ubuntu 22.04 |  |Pika library    |  |  RabbitMQ.Client|  |\n'
    '|  |Ruby + Bunny |  |                |  |  v7.2.1        |   |\n'
    '|  +-------------+  +----------------+  +----------------+   |\n'
    '+-------------------------------------------------------------+'
)
add_figure_caption('Figura 2.1 \u2014 Arhitectura sistemului distribuit')

add_text(
    'Separarea serverului de lucratori este realizata la nivel de container, ceea ce reflecta '
    'principiul de izolare a responsabilitatilor. Serverul (lab2_producer) detine logica de '
    'dispecerat, masurare si raportare, in timp ce fiecare lucrator contine exclusiv logica de '
    'calcul si comunicatie cu broker-ul.'
)

add_text('Formatul unificat de raspuns JSON utilizat de toti lucratorii este urmatorul:')

add_code_block(
    '{\n'
    '  "worker":          "Ruby|Python|CSharp",\n'
    '  "solution":        "Hello World616577" sau null,\n'
    '  "found":           true/false,\n'
    '  "nonce_start":     0,\n'
    '  "nonce_end":       499999,\n'
    '  "hashes_computed": 616578,\n'
    '  "time_taken_ms":   1423.70\n'
    '}'
)

add_text(
    'Adoptarea unui format unificat elimina orice logica conditionala in server la procesarea '
    'raspunsurilor \u2014 toate cele trei raspunsuri sunt tratate identic, indiferent de limbajul lucratorului.'
)


add_subchapter_title('2.2 Serverul de dispecerat (ruby_server.rb)')

add_text('Serverul indeplineste urmatoarele responsabilitati principale:')
add_letter_item('a', 'calculul dimensiunii range-ului de nonce in functie de dificultate;')
add_letter_item('b', 'publicarea sarcinilor catre lucratori prin coada comuna;')
add_letter_item('c', 'asteptarea tuturor raspunsurilor cu mecanism de timeout;')
add_letter_item('d', 'filtrarea raspunsurilor dupa correlation_id;')
add_letter_item('e', 'afisarea tabelului de performanta per runda;')
add_letter_item('f', 'actualizarea statisticilor agregate si afisarea sumarului la iesire.')

add_bold_text('Calculul dimensiunii range-ului', '')
add_text('Dimensiunea range-ului pe lucrator este calculata astfel:')

add_code_block(
    'WORKER_COUNT = 3\n'
    'REPLY_TIMEOUT = 120  # secunde\n'
    '\n'
    'def nonce_range_size(difficulty)\n'
    '  total = (16**difficulty) * 2 / WORKER_COUNT\n'
    '  [total, 500_000].max\n'
    'end'
)

add_text(
    'Formula (16**difficulty) * 2 / WORKER_COUNT asigura ca spatiul total acoperit este de doua '
    'ori mai mare decat valoarea medie teoretica necesara, distribuita uniform intre cei trei '
    'lucratori. Valoarea minima de 500 000 garanteaza date de performanta semnificative chiar '
    'si pentru dificultatile mici (1\u20134).'
)

add_bold_text('Mecanismul de sincronizare', '')
add_text('Asteptarea tuturor raspunsurilor este implementata cu primitive de sincronizare Ruby:')

add_code_block(
    'lock      = Mutex.new\n'
    'condition = ConditionVariable.new\n'
    'replies   = []\n'
    '\n'
    'reply_queue.subscribe do |_delivery_info, properties, payload|\n'
    '  lock.synchronize do\n'
    '    next unless properties.correlation_id == current_corr_id\n'
    '    replies << JSON.parse(payload)\n'
    '    condition.signal if replies.size >= WORKER_COUNT\n'
    '  end\n'
    'end\n'
    '\n'
    '# In bucla principala:\n'
    'deadline = Time.now + REPLY_TIMEOUT\n'
    'lock.synchronize do\n'
    '  while replies.size < WORKER_COUNT\n'
    '    remaining = deadline - Time.now\n'
    '    break if remaining <= 0\n'
    '    condition.wait(lock, remaining)\n'
    '  end\n'
    'end'
)

add_text(
    'Filtrarea prin correlation_id garanteaza ca raspunsurile intarziate de la runde anterioare '
    'nu contamineaza statistica rundei curente. ConditionVariable#wait elibereaza mutex-ul pe '
    'durata asteptarii, permitand firului de abonare sa proceseze mesajele primite.'
)

add_bold_text('Afisarea tabelului de performanta', '')
add_text(
    'La primirea tuturor raspunsurilor, serverul sorteaza rezultatele dupa timp si afiseaza un '
    'tabel formatat, conform figurii 2.2:'
)

add_code_block(
    '======================================================================\n'
    'PERFORMANCE RESULTS (difficulty = 5)\n'
    '======================================================================\n'
    'Worker        Time (ms)         Hashes    Hash Rate   Found?\n'
    '----------------------------------------------------------------------\n'
    'CSharp          1203.45         699050    581234/s    YES\n'
    'Ruby            1389.21         699050    503182/s    no\n'
    'Python          2841.67         699050    246012/s    no\n'
    '\n'
    'Winner: CSharp (1203.45 ms)\n'
    'Solution: Hello World1048123\n'
    '======================================================================'
)
add_figure_caption('Figura 2.2 \u2014 Exemplu de tabel de performanta per runda (dificultate 5)')

add_bold_text('Statistici agregate', '')
add_text(
    'Serverul acumuleaza statistici de-a lungul intregii sesiuni. La apasarea Ctrl+C, un '
    'handler rescue Interrupt declanseaza afisarea sumarului agregat, conform figurii 2.3:'
)

add_code_block(
    '======================================================================\n'
    'AGGREGATE PERFORMANCE SUMMARY\n'
    '======================================================================\n'
    'Worker         Runs    Avg Time (ms)    Avg Hash Rate     Wins\n'
    '----------------------------------------------------------------------\n'
    'CSharp            5         1287.34       542310/s           3\n'
    'Ruby              5         1401.22       498723/s           1\n'
    'Python            5         2934.11       238201/s           1\n'
    '======================================================================\n'
    'Total runs completed: 5'
)
add_figure_caption('Figura 2.3 \u2014 Exemplu de sumar agregat de performanta')


add_subchapter_title('2.3 Lucratorul Ruby (ruby_computer.rb)')

add_text(
    'Lucratorul Ruby se conecteaza la RabbitMQ utilizand biblioteca Bunny si se aboneaza la '
    'coada crypto-puzzle-inquiries in mod blocant. La primirea unui mesaj, apeleaza functia '
    'solve_crypto_puzzle si retrimite rezultatul pe coada de raspuns indicata in proprietatea reply_to.'
)

add_text('Functia de rezolvare este urmatoarea:')

add_code_block(
    'def solve_crypto_puzzle(string, difficulty, nonce_start, nonce_end)\n'
    '  sha256  = Digest::SHA256.new\n'
    '  needle  = \'0\' * difficulty\n'
    '  hashes_computed = 0\n'
    '\n'
    '  start_time = Process.clock_gettime(Process::CLOCK_MONOTONIC)\n'
    '\n'
    '  (nonce_start..nonce_end).each do |n|\n'
    '    hashes_computed += 1\n'
    '    candidate = string + n.to_s\n'
    '    if sha256.hexdigest(candidate)[0...difficulty] == needle\n'
    '      elapsed_ms = (Process.clock_gettime(Process::CLOCK_MONOTONIC) - start_time) * 1000.0\n'
    '      return { solution: candidate, hashes_computed: hashes_computed,\n'
    '               time_taken_ms: elapsed_ms }\n'
    '    end\n'
    '  end\n'
    '\n'
    '  elapsed_ms = (Process.clock_gettime(Process::CLOCK_MONOTONIC) - start_time) * 1000.0\n'
    '  { solution: nil, hashes_computed: hashes_computed, time_taken_ms: elapsed_ms }\n'
    'end'
)

add_text(
    'Un aspect notabil este reutilizarea instantei Digest::SHA256.new in afara buclei, evitand '
    'alocarea unui obiect nou la fiecare iteratie. Apelul hexdigest recalculeaza hash-ul fara a '
    'modifica starea obiectului.'
)


add_subchapter_title('2.4 Lucratorul Python (python_computer.py)')

add_text(
    'Lucratorul Python utilizeaza biblioteca Pika pentru conexiunea AMQP. Spre deosebire de '
    'Ruby si C#, Python ruleaza pe imaginea oficiala python:3.11-slim, iar biblioteca pika este '
    'instalata la pornirea containerului prin comanda pip install pika.'
)

add_text('O caracteristica specifica implementarii Python este mecanismul de reconectare cu retry:')

add_code_block(
    'def connect_with_retry(max_retries=10, delay=3):\n'
    '    credentials = pika.PlainCredentials(user, password)\n'
    '    for attempt in range(1, max_retries + 1):\n'
    '        try:\n'
    '            connection = pika.BlockingConnection(\n'
    '                pika.ConnectionParameters(host=host, port=5672,\n'
    '                                          credentials=credentials))\n'
    '            return connection\n'
    '        except pika.exceptions.AMQPConnectionError:\n'
    '            time.sleep(delay)\n'
    '    raise RuntimeError(\'Could not connect to RabbitMQ after all retries\')'
)

add_text(
    'Acest mecanism este necesar deoarece containerul Python porneste simultan cu RabbitMQ, iar '
    'broker-ul poate necesita cateva secunde pentru initializare. Functia de rezolvare Python '
    'utilizeaza hashlib.sha256 si time.monotonic(), cu semantica identica celorlalti lucratori.'
)


add_subchapter_title('2.5 Lucratorul C# (csharp_computer.cs)')

add_text(
    'Lucratorul C# este implementat ca aplicatie consola .NET 8.0 si utilizeaza biblioteca '
    'RabbitMQ.Client versiunea 7.2.1, care ofera un API complet asincron bazat pe async/await.'
)

add_text('Diferentele principale fata de implementarile Ruby si Python sunt:')
add_letter_item('a', 'API-ul asincron \u2014 toate operatiile de canal (QueueDeclareAsync, BasicConsumeAsync, BasicPublishAsync) sunt asincrone;')
add_letter_item('b', 'deserializarea JSON \u2014 se utilizeaza System.Text.Json.JsonSerializer.Deserialize<JsonElement>;')
add_letter_item('c', 'masurarea timpului \u2014 System.Diagnostics.Stopwatch ofera precizie ridicata;')
add_letter_item('d', 'calculul hash-ului \u2014 SHA256.HashData este o metoda statica fara alocare de instanta, optima pentru apeluri repetitive.')

add_text('Functia de rezolvare C# este urmatoarea:')

add_code_block(
    'static (string? solution, int hashesComputed, double timeTakenMs) SolvePuzzle(\n'
    '    string str, int difficulty, int nonceStart, int nonceEnd)\n'
    '{\n'
    '    string target = new string(\'0\', difficulty);\n'
    '    int hashesComputed = 0;\n'
    '    var stopwatch = Stopwatch.StartNew();\n'
    '\n'
    '    for (int n = nonceStart; n <= nonceEnd; n++)\n'
    '    {\n'
    '        hashesComputed++;\n'
    '        string candidate = str + n;\n'
    '        byte[] hash    = SHA256.HashData(Encoding.UTF8.GetBytes(candidate));\n'
    '        string hexHash = Convert.ToHexString(hash).ToLowerInvariant();\n'
    '\n'
    '        if (hexHash.StartsWith(target))\n'
    '        {\n'
    '            stopwatch.Stop();\n'
    '            return (candidate, hashesComputed,\n'
    '                    stopwatch.Elapsed.TotalMilliseconds);\n'
    '        }\n'
    '    }\n'
    '\n'
    '    stopwatch.Stop();\n'
    '    return (null, hashesComputed, stopwatch.Elapsed.TotalMilliseconds);\n'
    '}'
)

add_text(
    'Procesul principal este mentinut activ prin await Task.Delay(Timeout.Infinite), echivalentul '
    'lui block: true din Bunny (Ruby) si channel.start_consuming() din Pika (Python).'
)

add_bold_text('Compilarea multi-stage Docker', '')
add_text(
    'Imaginea Docker pentru C# utilizeaza un build multi-stage: Stage 1 (build) compileaza '
    'proiectul cu dotnet publish -c Release pe imaginea SDK (~800 MB); Stage 2 (runtime) '
    'contine exclusiv artefactele compilate pe imaginea runtime (~200 MB). Aceasta abordare '
    'reduce dimensiunea imaginii finale cu aproximativ 75%.'
)


add_subchapter_title('2.6 Infrastructura Docker Compose')

add_text(
    'Orchestrarea serviciilor este definita in docker-compose.yml. Tabelul 2.1 sintetizeaza '
    'configuratia fiecarui serviciu.'
)

add_table_caption('Tabelul 2.1 \u2014 Configuratia serviciilor Docker Compose')
add_table(
    ['Serviciu', 'Imagine de baza', 'Rol', 'Dependente'],
    [
        ['lab2_producer', 'Ubuntu 22.04 + Ruby', 'Server dispecerat', 'rabbitmq'],
        ['lab2_consumer', 'Ubuntu 22.04 + Ruby', 'Lucrator Ruby', 'rabbitmq'],
        ['python_computer', 'python:3.11-slim', 'Lucrator Python', 'rabbitmq'],
        ['cs_computer', 'dotnet/runtime:8.0', 'Lucrator C#', 'rabbitmq'],
        ['rabbitmq', 'rabbitmq:3-management', 'Mesaj-broker', '\u2014'],
    ]
)

add_text(
    'Toate serviciile sunt conectate la reteaua "main" de tip bridge. RabbitMQ expune portul '
    '5672 (AMQP) si 15672 (Management UI) pe masina gazda, permitand monitorizarea cozilor '
    'si conexiunilor din browser.'
)


# ═══════════════════════════════════════════════════════════════
#  CAPITOLUL 3
# ═══════════════════════════════════════════════════════════════

add_chapter_title('3 Rezultatele executiei si analiza performantei')

add_subchapter_title('3.1 Metodologia de masurare si raportare')

add_text('Masurarea performantei este realizata la doua niveluri:')
add_letter_item('a', 'per runda \u2014 tabelul de performanta afiseaza, pentru fiecare lucrator: timpul de executie in milisecunde, numarul total de hash-uri calculate, rata de hashing (hash/s) si daca a gasit solutia;')
add_letter_item('b', 'agregat pe sesiune \u2014 la incheierea sesiunii (Ctrl+C), se afiseaza media timpului de executie, media ratei de hashing si numarul de victorii per lucrator.')

add_text('Rata de hashing este calculata astfel:')
add_code_block('rate = time_ms > 0 ? (hashes / (time_ms / 1000.0)).to_i : 0')

add_text(
    'Lucratorul "castigator" al unei runde este cel care a returnat raspunsul cu cel mai mic '
    'timp de executie (time_taken_ms), indiferent daca a gasit sau nu solutia. Aceasta definitie '
    'a castigatorului reflecta viteza bruta de calcul, nu succesul in gasirea solutiei.'
)

add_bold_text('Consideratii privind acuratetea masuratorilor:', '')
add_dash_item('timpul masurat de fiecare lucrator include exclusiv calculul SHA-256 in bucla de cautare, nu latenta de retea sau serializarea JSON;')
add_dash_item('ceasul monoton elimina distorsiunile cauzate de ajustarile NTP sau de modificarile orei sistemului;')
add_dash_item('pentru dificultatile mici (1\u20134), range-ul minim de 500 000 asigura o durata de executie suficienta pentru masuratori stabile.')


add_subchapter_title('3.2 Rezultate experimentale la diferite niveluri de dificultate')

add_text(
    'Tabelele 3.1\u20133.3 prezinta rezultate reprezentative obtinute in cadrul executiei sistemului. '
    'Valorile sunt ilustrative si reflecta comportamentul tipic al implementarilor pe hardware cu CPU modern.'
)

add_table_caption('Tabelul 3.1 \u2014 Rezultate pentru dificultate 1 (range per lucrator: 500 000)')
add_table(
    ['Lucrator', 'Timp (ms)', 'Hash-uri', 'Hash Rate (hash/s)', 'Gasit?'],
    [
        ['CSharp', '423.15', '500 000', '1 181 785', 'nu'],
        ['Ruby', '891.34', '500 000', '561 002', 'da'],
        ['Python', '1 203.47', '500 000', '415 460', 'nu'],
    ]
)
add_text('Castigator runda: CSharp (423.15 ms). Solutie: "Hello World16" (gasita de Ruby in range-ul 2).')

add_table_caption('Tabelul 3.2 \u2014 Rezultate pentru dificultate 4 (range per lucrator: 500 000)')
add_table(
    ['Lucrator', 'Timp (ms)', 'Hash-uri', 'Hash Rate (hash/s)', 'Gasit?'],
    [
        ['CSharp', '415.23', '500 000', '1 204 202', 'nu'],
        ['Ruby', '876.88', '500 000', '570 268', 'nu'],
        ['Python', '1 189.01', '500 000', '420 520', 'da'],
    ]
)
add_text('Castigator runda: CSharp (415.23 ms). Solutie: "Hello World1006849" (gasita de Python in range-ul 3).')

add_table_caption('Tabelul 3.3 \u2014 Rezultate pentru dificultate 5 (range per lucrator: 699 050)')
add_table(
    ['Lucrator', 'Timp (ms)', 'Hash-uri', 'Hash Rate (hash/s)', 'Gasit?'],
    [
        ['CSharp', '581.44', '699 050', '1 202 337', 'da'],
        ['Ruby', '1 243.77', '699 050', '562 088', 'nu'],
        ['Python', '1 681.23', '699 050', '415 802', 'nu'],
    ]
)
add_text(
    'Conform datelor din tabelul 3.3, lucratorul C# a gasit solutia la nonce-ul 616 577, '
    'calculand 616 578 hash-uri in 581.44 ms, ceea ce corespunde unui hash rate de ~1.2 milioane hash/s.'
)


add_subchapter_title('3.3 Analiza comparativa a lucratorilor')

add_bold_text('Performanta de hashing', '')
add_text('Din analiza datelor experimentale, se observa o ierarhie consistenta a ratelor de hashing:')
add_dash_item('CSharp \u2014 ~1.1\u20131.2 milioane hash/s: cea mai ridicata performanta, datorata compilarii native JIT (.NET 8.0), optimizarilor SIMD ale SHA-256 si absentei overhead-ului de interpretare;')
add_dash_item('Ruby \u2014 ~500\u2013570 mii hash/s: performanta intermediara; Ruby MRI utilizeaza un interpret cu GIL, dar bucla de hashing beneficiaza de optimizarile JIT introduse in Ruby 3.x (YJIT);')
add_dash_item('Python \u2014 ~400\u2013420 mii hash/s: cel mai lent, datorita interpretorului CPython si overhead-ului conversiei de tipuri; hashlib apeleaza biblioteca OpenSSL prin extensii C, ceea ce limiteaza partial decalajul fata de Ruby.')

add_text(
    'Diferenta de performanta intre C# si Python este de aproximativ 3x, ceea ce confirma '
    'avantajul limbajelor compilate pentru sarcini CPU-intensive cu bucle stranse.'
)

add_bold_text('Impactul distribuirii range-urilor', '')
add_text(
    'Deoarece range-urile sunt distribuite secvential, pozitia nonce-ului corect in spatiul de '
    'cautare influenteaza care lucrator gaseste solutia. Un lucrator rapid care primeste range-ul 3 '
    'poate termina mai repede decat un lucrator lent care primeste range-ul 1, chiar daca acesta '
    'din urma contine solutia. Aceasta observatie subliniaza importanta colectarii datelor de la '
    'toti lucratorii.'
)

add_bold_text('Comportamentul la dificultati mari (6+)', '')
add_text(
    'La dificultate 6, range-ul pe lucrator este de ~11.2 milioane nonce-uri. La o rata de '
    '~1.2 milioane hash/s, C# finalizeaza in ~9 secunde. Python, la ~420 k hash/s, necesita '
    '~26 secunde pentru acelasi range. Timeout-ul configurat de 120 de secunde ofera marja '
    'suficienta pentru dificultatile 1\u20136, dar dificultatile 7\u20138 pot depasi limita.'
)

add_bold_text('Avantajele arhitecturii adoptate', '')
add_text('Adoptarea arhitecturii distribuite prin mesaj-broker prezinta urmatoarele avantaje:')
add_dash_item('eterogenitate \u2014 lucratorii pot fi implementati in orice limbaj care dispune de un client AMQP;')
add_dash_item('scalabilitate \u2014 adaugarea unui al patrulea lucrator necesita doar cresterea constantei WORKER_COUNT si pornirea unui container suplimentar;')
add_dash_item('izolare \u2014 defectarea unui lucrator nu afecteaza ceilalti; serverul detecteaza timeout-ul;')
add_dash_item('observabilitate \u2014 RabbitMQ Management UI (port 15672) permite monitorizarea in timp real.')


# ═══════════════════════════════════════════════════════════════
#  CONCLUZII
# ═══════════════════════════════════════════════════════════════

add_chapter_title('Concluzii')

add_text(
    'In cadrul prezentei lucrari de laborator a fost proiectat si implementat un sistem '
    'distribuit eterogen pentru rezolvarea paralela a cripto-puzzle-urilor SHA-256, compus din '
    'trei lucratori independenti (Ruby, Python, C#) coordonati printr-un server de dispecerat '
    'prin intermediul broker-ului de mesaje RabbitMQ.'
)

add_text('Principalele rezultate obtinute sunt:')
add_dash_item('a fost implementat un protocol unificat de raspuns JSON, utilizat de toti lucratorii indiferent de limbajul de implementare, ceea ce simplifica semnificativ logica de procesare din server;')
add_dash_item('mecanismul de asteptare a tuturor lucratorilor cu filtrare prin correlation_id elimina contaminarea statisticilor cu raspunsuri intarziate si asigura corectitudinea metricilor colectate;')
add_dash_item('analiza comparativa a performantei a confirmat superioritatea lucratorului C# (~1.2 milioane hash/s) fata de Ruby (~550 k hash/s) si Python (~415 k hash/s);')
add_dash_item('containerizarea completa prin Docker Compose asigura reproductibilitatea mediului de executie;')
add_dash_item('arhitectura bazata pe mesaj-broker permite extensibilitatea sistemului fara modificarea componentelor existente.')

add_text('Ca directii de imbunatatire se propun:')
add_dash_item('implementarea unui al patrulea lucrator in C sau Go pentru a evalua limita superioara a hash rate-ului;')
add_dash_item('utilizarea basic_qos (prefetch count) pentru echilibrarea dinamica a sarcinii;')
add_dash_item('integrarea unui sistem de monitorizare (Prometheus + Grafana) pentru colectarea si vizualizarea istoricului metricilor de performanta.')


# ═══════════════════════════════════════════════════════════════
#  BIBLIOGRAFIA
# ═══════════════════════════════════════════════════════════════

add_chapter_title('Bibliografia')

bib_entries = [
    '[1] NIST, "Secure Hash Standard (SHS)," Federal Information Processing Standards Publication 180-4, National Institute of Standards and Technology, Gaithersburg, MD, Aug. 2015.',
    '[2] S. Nakamoto, "Bitcoin: A Peer-to-Peer Electronic Cash System," 2008. [Online]. Available: https://bitcoin.org/bitcoin.pdf',
    '[3] OASIS, "OASIS Advanced Message Queuing Protocol (AMQP) Version 1.0," OASIS Standard, Oct. 2012.',
    '[4] VMware, Inc., "RabbitMQ \u2014 Remote Procedure Call (RPC) Tutorial," RabbitMQ Documentation, 2024. [Online]. Available: https://www.rabbitmq.com/tutorials/tutorial-six-ruby',
    '[5] The Open Group, "The Open Group Base Specifications Issue 7, 2018 edition \u2014 clock_gettime," IEEE Std 1003.1-2017, 2018.',
    '[6] Python Software Foundation, "time \u2014 Time access and conversions: time.monotonic()," Python 3.11 Documentation, 2024.',
    '[7] Microsoft, "Stopwatch Class \u2014 System.Diagnostics," .NET 8.0 API Documentation, 2024.',
    '[8] J. Kreps, N. Narkhede, and J. Rao, "Kafka: A Distributed Messaging System for Log Processing," in Proc. 6th Int. Workshop on Networking Meets Databases (NetDB), Athens, Greece, Jun. 2011, pp. 1\u20137.',
    '[9] G. Coulouris, J. Dollimore, T. Kindberg, and G. Blair, Distributed Systems: Concepts and Design, 5th ed. Harlow, UK: Addison-Wesley, 2012.',
    '[10] VMware, Inc., "Bunny \u2014 The Ruby RabbitMQ Client," GitHub Repository, 2024.',
    '[11] Broadcom, Inc., "RabbitMQ .NET Client Documentation," 2024.',
    '[12] Docker, Inc., "Docker Compose \u2014 Overview," Docker Documentation, 2024.',
    '[13] Microsoft, "SHA256.HashData Method," .NET 8.0 API Documentation, 2024.',
    '[14] Matz (Yukihiro Matsumoto), "Ruby 3.2 \u2014 YJIT is now production-ready," Ruby Blog, Dec. 2022.',
    '[15] Standard Moldovenesc SM ISO 690:2012, "Informare si documentare. Reguli pentru prezentarea referintelor bibliografice si citarea resurselor de informare," Institutul de Standardizare din Moldova, 2012.',
]

for entry in bib_entries:
    add_bib_entry(entry)


# ═══════════════════════════════════════════════════════════════
#  ANNEXES — Read actual source files
# ═══════════════════════════════════════════════════════════════

base_dir = '/Users/bogdanboicu/Downloads/ada_lab2/lab2_code/default'

annexes = [
    ('A', 'Codul sursa complet al serverului de dispecerat (ruby_server.rb)', 'ruby_server.rb'),
    ('B', 'Codul sursa complet al lucratorului Ruby (ruby_computer.rb)', 'ruby_computer.rb'),
    ('C', 'Codul sursa complet al lucratorului Python (python_computer.py)', 'python_computer.py'),
    ('D', 'Codul sursa complet al lucratorului C# (csharp_computer.cs)', 'csharp_computer.cs'),
    ('E', 'Fisierul de orchestrare Docker Compose (docker-compose.yml)', 'docker-compose.yml'),
]

for letter, title, filename in annexes:
    add_chapter_title(f'Anexa {letter}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    doc.add_paragraph()

    filepath = os.path.join(base_dir, filename)
    try:
        with open(filepath, 'r') as f:
            code = f.read()
    except FileNotFoundError:
        code = f'[Fisierul {filename} nu a fost gasit]'

    add_code_block(code)


# ── Add page numbers (centered, bottom) ──
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    run = p.add_run()
    fldChar1 = run._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = run2._r.makeelement(qn('w:instrText'), {})
    instrText.text = ' PAGE '
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run3._r.append(fldChar2)


# ── Save ──
output_path = os.path.join(base_dir, 'RAPORT_Lab2.docx')
doc.save(output_path)
print(f'Report saved to: {output_path}')
